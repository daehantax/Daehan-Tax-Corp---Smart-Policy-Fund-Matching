# =============================================================================
# [조사용] C-1 첨부파일 실태조사 — 2단계: 문서 텍스트 추출 + 통계 집계
#
# ⚠ 조사(research) 전용. 운영 파이프라인에 연결하지 않는다. → docs/WORK_PLAN.md C-1
#
# 사용법:
#   py -m pip install --target .research-data/pylibs -r scripts/research/requirements.txt
#   PYTHONPATH=.research-data/pylibs py scripts/research/extract_survey.py
#
# 입력:  .research-data/manifest.json  (fetch-attachments.mjs 산출물)
# 출력:  .research-data/survey-results.json   추출 결과 원본
#        docs/research/attachment-survey.md   사람이 읽는 리포트
#
# 측정 항목 (WORK_PLAN C-1 체크리스트와 1:1 대응):
#   확장자 분포 / 공고당 첨부 개수 / 파일명 규칙 / 추출 성공률 /
#   글자수 분포 / 스캔 PDF 비율 / 표 포함 비율 / 「제외 대상」 문구 패턴
# =============================================================================

import io
import json
import os
import re
import struct
import sys
import zipfile
import zlib
from collections import Counter, defaultdict
from pathlib import Path

ROOT = Path(__file__).resolve().parents[2]
DATA = Path(os.environ.get("SURVEY_OUT", ROOT / ".research-data"))
MANIFEST = DATA / "manifest.json"
FILES = DATA / "files"
OUT_JSON = DATA / "survey-results.json"
OUT_MD = ROOT / "docs" / "research" / "attachment-survey.md"

# 자격 조건이 담긴 문서인지 판별하는 신호
ELIGIBILITY_CUES = [
    "지원대상", "신청자격", "지원자격", "신청대상", "참여자격",
    "제외대상", "제외 대상", "지원제외", "신청제외",
]
# 「신청 제외 대상」 문구 패턴 수집용
EXCLUSION_RE = re.compile(r"(?:신청\s*)?제외\s*대상|지원\s*제외|참여\s*제한|신청\s*제한")
# 우리가 판정하지 못하는 수치 조건 (WORK_PLAN M-4 대상)
NUMERIC_CUES = {
    "업력": re.compile(r"업력|창업\s*\d+\s*년|설립\s*\d+\s*년"),
    "매출액": re.compile(r"매출액|연\s*매출|평균\s*매출"),
    "근로자수": re.compile(r"상시\s*근로자|근로자\s*\d+\s*인|종업원\s*\d+"),
    "업종코드": re.compile(r"한국표준산업분류|표준산업분류|KSIC|업종코드"),
    "인증": re.compile(r"벤처기업\s*확인|이노비즈|메인비즈|인증\s*기업"),
}


# ── 형식별 추출기 ─────────────────────────────────────────────────────────────

def extract_hwp(path):
    """HWP 5.0 = OLE 복합문서. BodyText/Section* 을 풀고 문단 텍스트 레코드만 읽는다."""
    import olefile

    f = olefile.OleFileIO(str(path))
    names = ["/".join(s) for s in f.listdir()]
    if "FileHeader" not in names:
        raise ValueError("FileHeader 없음 — HWP 5.0 아님")
    compressed = bool(f.openstream("FileHeader").read()[36] & 1)

    parts, tables = [], 0
    for name in sorted(n for n in names if n.startswith("BodyText")):
        d = f.openstream(name).read()
        if compressed:
            d = zlib.decompress(d, -15)
        i = 0
        while i + 4 <= len(d):
            header = struct.unpack("<I", d[i:i + 4])[0]
            tag, size = header & 0x3FF, (header >> 20) & 0xFFF
            i += 4
            if size == 0xFFF:
                size = struct.unpack("<I", d[i:i + 4])[0]
                i += 4
            if tag == 76:            # HWPTAG_TABLE
                tables += 1
            elif tag == 67:          # HWPTAG_PARA_TEXT
                # 본문은 UTF-16LE 다. 코드 단위를 chr() 로 하나씩 만들면 서로게이트 쌍이
                # 깨져(lone surrogate) 나중에 UTF-8 로 저장할 때 터진다.
                # 그래서 바이트로 모아 두고 마지막에 한 번에 디코드한다.
                raw, buf, j = d[i:i + size], bytearray(), 0
                while j + 1 < len(raw):
                    c = struct.unpack("<H", raw[j:j + 2])[0]
                    if c in (0, 10, 13):
                        buf += "\n".encode("utf-16-le"); j += 2
                    elif c < 32:
                        # 확장/인라인 컨트롤은 16바이트, 나머지는 2바이트
                        j += 16 if c in (1, 2, 3, 11, 12, 14, 15, 16, 17, 18, 21, 22, 23) else 2
                    else:
                        buf += raw[j:j + 2]; j += 2
                parts.append(buf.decode("utf-16-le", "replace"))
            i += size
    f.close()
    return "\n".join(parts), {"tables": tables}


def extract_hwpx(path):
    """HWPX = ZIP + XML. Contents/section*.xml 의 텍스트 노드를 모은다."""
    text, tables = [], 0
    with zipfile.ZipFile(path) as z:
        for name in sorted(n for n in z.namelist() if re.match(r"Contents/section\d+\.xml", n)):
            xml = z.read(name).decode("utf-8", "replace")
            tables += len(re.findall(r"<hp:tbl[\s>]", xml))
            # <hp:t>본문</hp:t> — 네임스페이스 접두는 파일마다 다를 수 있어 느슨하게 잡는다
            for m in re.finditer(r"<(?:\w+:)?t(?:\s[^>]*)?>(.*?)</(?:\w+:)?t>", xml, re.S):
                chunk = re.sub(r"<[^>]+>", "", m.group(1))
                text.append(chunk.replace("&amp;", "&").replace("&lt;", "<").replace("&gt;", ">"))
    return "\n".join(text), {"tables": tables}


def extract_pdf(path):
    """PDF. 페이지별 텍스트와 표를 세고, 텍스트가 거의 없으면 스캔본으로 본다."""
    import pdfplumber

    parts, tables, pages, empty_pages = [], 0, 0, 0
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            pages += 1
            t = page.extract_text() or ""
            if len(t.strip()) < 20:
                empty_pages += 1
            parts.append(t)
            try:
                tables += len(page.find_tables())
            except Exception:
                pass
    text = "\n".join(parts)
    scanned = pages > 0 and empty_pages / pages >= 0.8
    return text, {"tables": tables, "pages": pages, "scanned": scanned}


def extract_xlsx(path):
    import openpyxl

    wb = openpyxl.load_workbook(str(path), data_only=True, read_only=True)
    parts = []
    for ws in wb.worksheets:
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None]
            if cells:
                parts.append(" ".join(cells))
    wb.close()
    # 스프레드시트는 그 자체가 표다
    return "\n".join(parts), {"tables": len(parts) and 1}


def extract_docx(path):
    import docx

    d = docx.Document(str(path))
    parts = [p.text for p in d.paragraphs]
    for t in d.tables:
        for row in t.rows:
            parts.append(" ".join(c.text for c in row.cells))
    return "\n".join(parts), {"tables": len(d.tables)}


def extract_zip(path):
    """ZIP 은 내용을 풀지 않고 안에 무엇이 들었는지만 기록한다 (A-3 설계 판단용)."""
    with zipfile.ZipFile(path) as z:
        inner = [n for n in z.namelist() if not n.endswith("/")]
    exts = Counter(Path(n).suffix.lower().lstrip(".") or "(없음)" for n in inner)
    return "", {"zip_entries": len(inner), "zip_exts": dict(exts)}


EXTRACTORS = {
    "hwp": extract_hwp, "hwpx": extract_hwpx, "pdf": extract_pdf,
    "xlsx": extract_xlsx, "xlsm": extract_xlsx, "docx": extract_docx, "zip": extract_zip,
}
# 문서가 아니라 추출 대상이 아닌 것들 — 실패가 아니라 '해당없음'으로 센다
NON_DOC = {"jpg", "jpeg", "png", "gif", "bmp", "tif", "tiff", "mp4", "avi", "hwt", "ppt", "pptx"}


def sniff_format(path, ext):
    """실제 형식을 매직 바이트로 판별한다.

    확장자를 믿으면 안 된다 — 표본에서 `.hwpx` 로 올라온 파일 2개가 실제로는
    HWP 5.0(OLE)이었다. 확장자만 보고 zipfile 로 열면 BadZipFile 로 죽는다.
    """
    with open(path, "rb") as f:
        head = f.read(8)
    if head.startswith(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"):
        return "hwp"            # OLE 복합문서 — HWP 5.0 (드물게 xls/doc)
    if head.startswith(b"%PDF"):
        return "pdf"
    if head.startswith(b"PK\x03\x04"):
        # ZIP 계열 — hwpx / xlsx / docx / 진짜 zip 은 확장자로 가른다
        return ext if ext in ("hwpx", "xlsx", "xlsm", "docx", "zip") else "zip"
    return ext


# ── 파일명 분류 ───────────────────────────────────────────────────────────────

def classify_name(name):
    """파일명으로 본문(공고문)인지 서식인지 가른다. A-2 선별 규칙의 근거가 된다."""
    n = name.replace(" ", "")
    if re.search(r"공고문|공고$|모집공고|공고\)|사업공고", n):
        return "공고문"
    if re.search(r"신청서|지원서|참가신청|접수서", n):
        return "신청서"
    if re.search(r"동의서|개인정보", n):
        return "동의서"
    if re.search(r"서식|양식|별지", n):
        return "서식"
    if re.search(r"계획서|사업계획", n):
        return "계획서"
    if re.search(r"안내|매뉴얼|가이드|설명", n):
        return "안내문"
    if re.search(r"^\[?붙임|^\(?붙임|첨부\d", n):
        return "붙임(기타)"
    return "기타"


# ── 집계 ─────────────────────────────────────────────────────────────────────

def pct(a, b):
    return f"{(100 * a / b):.1f}%" if b else "—"


def main():
    if not MANIFEST.exists():
        sys.exit(f"[survey] manifest 없음: {MANIFEST}\n먼저 node scripts/research/fetch-attachments.mjs 를 실행하세요.")

    manifest = json.loads(MANIFEST.read_text(encoding="utf-8"))
    results = []

    for entry in manifest:
        for att in entry["attachments"]:
            ext = att["ext"]
            rec = {
                "pblancId": entry["pblancId"],
                "title": entry["title"],
                "fileName": att["fileName"],
                "ext": ext,
                "bytes": att.get("bytes"),
                "nameClass": classify_name(att["fileName"]),
                "status": None, "chars": 0, "tables": 0,
                "scanned": False, "hasEligibility": False,
                "numericCues": [], "exclusionHits": [],
                "error": None,
            }
            path = FILES / att["localName"]

            real = sniff_format(path, ext) if path.exists() else ext
            rec["realFormat"] = real
            rec["extMismatch"] = real != ext

            if att["status"] not in ("ok", "cached") or not path.exists():
                rec["status"] = "다운로드실패"
            elif real in NON_DOC:
                rec["status"] = "문서아님"
            elif real not in EXTRACTORS:
                rec["status"] = "미지원형식"
            else:
                try:
                    text, meta = EXTRACTORS[real](path)
                    text = re.sub(r"[ \t]+", " ", text)
                    rec["status"] = "성공" if len(text.strip()) > 0 or ext == "zip" else "빈텍스트"
                    rec["chars"] = len(text.strip())
                    rec["tables"] = meta.get("tables", 0) or 0
                    rec["scanned"] = bool(meta.get("scanned"))
                    rec["zip"] = meta.get("zip_exts")
                    if text:
                        rec["hasEligibility"] = any(c in text for c in ELIGIBILITY_CUES)
                        rec["numericCues"] = [k for k, r in NUMERIC_CUES.items() if r.search(text)]
                        rec["exclusionHits"] = list({m.group(0) for m in EXCLUSION_RE.finditer(text)})
                        (DATA / "text").mkdir(exist_ok=True)
                        (DATA / "text" / (att["localName"] + ".txt")).write_text(text, encoding="utf-8")
                except Exception as e:
                    rec["status"] = "추출실패"
                    rec["error"] = f"{type(e).__name__}: {e}"
            results.append(rec)

    OUT_JSON.write_text(json.dumps(results, ensure_ascii=False, indent=2), encoding="utf-8")

    # ── 리포트 ────────────────────────────────────────────────────────────────
    n_prog = len(manifest)
    n_att = len(results)
    per_prog = Counter(len(e["attachments"]) for e in manifest)
    ext_count = Counter(r["ext"] for r in results)
    name_count = Counter(r["nameClass"] for r in results)

    by_ext = defaultdict(lambda: {"total": 0, "ok": 0, "chars": []})
    for r in results:
        b = by_ext[r["ext"]]
        b["total"] += 1
        if r["status"] == "성공":
            b["ok"] += 1
            b["chars"].append(r["chars"])

    docs = [r for r in results if r["status"] == "성공"]
    chars = sorted(r["chars"] for r in docs) or [0]
    pdfs = [r for r in results if r.get("realFormat") == "pdf" and r["status"] in ("성공", "빈텍스트")]
    scanned = [r for r in pdfs if r["scanned"]]
    with_tables = [r for r in docs if r["tables"] > 0]
    elig = [r for r in docs if r["hasEligibility"]]

    # 자격 조건을 담은 파일의 파일명 분류 — A-2 선별 규칙의 핵심 근거
    elig_by_class = Counter(r["nameClass"] for r in elig)
    class_totals = Counter(r["nameClass"] for r in docs)

    # 공고 단위: 자격 조건이 담긴 첨부가 하나라도 있는가
    prog_has_elig = {r["pblancId"] for r in elig}
    numeric_tally = Counter(c for r in docs for c in r["numericCues"])
    numeric_prog = defaultdict(set)
    for r in docs:
        for c in r["numericCues"]:
            numeric_prog[c].add(r["pblancId"])

    L = []
    w = L.append
    w("# C-1 첨부파일 실태조사 결과\n")
    w("> `scripts/research/fetch-attachments.mjs` + `extract_survey.py` 산출물.")
    w("> 원본 데이터: `.research-data/survey-results.json` (커밋하지 않음)\n")
    w(f"- 표본 공고 **{n_prog}건**, 첨부파일 **{n_att}개**")
    w(f"- 공고당 평균 첨부 **{n_att / n_prog:.1f}개**\n")

    w("## 1. 확장자 분포\n")
    w("| 확장자 | 개수 | 비율 | 추출 성공 | 성공률 |")
    w("|---|---:|---:|---:|---:|")
    for ext, cnt in ext_count.most_common():
        b = by_ext[ext]
        w(f"| `{ext}` | {cnt} | {pct(cnt, n_att)} | {b['ok']} | {pct(b['ok'], b['total'])} |")

    w("\n## 2. 공고당 첨부 개수\n")
    w("| 첨부 수 | 공고 수 |")
    w("|---:|---:|")
    for k in sorted(per_prog):
        w(f"| {k} | {per_prog[k]} |")

    w("\n## 3. 파일명 분류\n")
    w("자격 조건이 어느 종류의 파일에 들어 있는지 — **A-2 선별 규칙의 근거**\n")
    w("| 분류 | 개수 | 추출 성공 | 자격조건 포함 | 포함률 |")
    w("|---|---:|---:|---:|---:|")
    for cls, cnt in name_count.most_common():
        ok = class_totals.get(cls, 0)
        e = elig_by_class.get(cls, 0)
        w(f"| {cls} | {cnt} | {ok} | {e} | {pct(e, ok)} |")

    w("\n### 3-1. 선별 전략 비교\n")
    w("파일명으로 「공고문」만 고르면 자격 조건을 놓치는 공고가 생긴다.\n")
    w("| 전략 | 자격조건 확보 공고 | AI 투입 문서 | 투입 글자수 |")
    w("|---|---:|---:|---:|")
    elig_chars = sum(r["chars"] for r in elig)
    gongo = [r for r in docs if r["nameClass"] == "공고문"]
    gongo_prog = {r["pblancId"] for r in gongo if r["hasEligibility"]}
    w(f"| 전부 투입 | {len(prog_has_elig)} | {len(docs)} | {sum(chars):,}자 |")
    w(f"| 내용으로 선별 (자격 단서 있는 문서만) | {len(prog_has_elig)} | {len(elig)} | {elig_chars:,}자 |")
    w(f"| 파일명으로 선별 (「공고문」만) | {len(gongo_prog)} | {len(gongo)} | {sum(r['chars'] for r in gongo):,}자 |")
    w("")
    w(f"→ **파일명 선별은 {len(prog_has_elig) - len(gongo_prog)}건을 놓친다.** "
      f"추출은 싸고 AI가 비싸므로 **전부 추출한 뒤 내용으로 거르는 쪽**이 맞다.")

    w("\n## 4. 추출 텍스트 분량\n")
    if docs:
        w(f"- 추출 성공 **{len(docs)}개**")
        w(f"- 글자수 — 중앙값 **{chars[len(chars) // 2]:,}자** / 평균 **{sum(chars) // len(chars):,}자** / 최대 **{max(chars):,}자**")
        w(f"- 합계 **{sum(chars):,}자**")
        w(f"- 내용 선별 후 AI 투입 분량 **{elig_chars:,}자** (공고당 평균 **{elig_chars // n_prog:,}자**) → C-6 입력값")
    else:
        w("- 추출 성공 0개")

    mism = [r for r in results if r.get("extMismatch")]
    w("\n### 4-1. 확장자 ≠ 실제 형식\n")
    if mism:
        w(f"**{len(mism)}건**. 확장자를 믿고 파서를 고르면 죽는다 — 매직 바이트로 판별해야 한다.\n")
        w("| 파일명 | 확장자 | 실제 |")
        w("|---|---|---|")
        for r in mism[:15]:
            w(f"| {r['fileName'][:45]} | `{r['ext']}` | `{r['realFormat']}` |")
    else:
        w("없음.")

    w("\n## 5. 리스크 지표\n")
    w("| 항목 | 값 | 뜻 |")
    w("|---|---:|---|")
    w(f"| 스캔 이미지 PDF | {len(scanned)} / {len(pdfs)} ({pct(len(scanned), len(pdfs))}) | OCR 없이는 포기 구간 |")
    w(f"| 표 포함 문서 | {len(with_tables)} / {len(docs)} ({pct(len(with_tables), len(docs))}) | 평문 추출 시 구조 손실 |")
    w(f"| 자격조건 포함 문서 | {len(elig)} / {len(docs)} ({pct(len(elig), len(docs))}) | 전부 AI에 넣을 필요 없음 |")
    w(f"| 자격조건 확보 공고 | {len(prog_has_elig)} / {n_prog} ({pct(len(prog_has_elig), n_prog)}) | 첨부로 조건을 읽을 수 있는 비율 |")

    fails = [r for r in results if r["status"] in ("추출실패", "빈텍스트", "미지원형식", "다운로드실패")]
    if fails:
        w(f"\n### 실패·미지원 {len(fails)}건\n")
        w("| 파일 | 확장자 | 상태 | 사유 |")
        w("|---|---|---|---|")
        for r in fails[:25]:
            w(f"| {r['fileName'][:40]} | `{r['ext']}` | {r['status']} | {(r['error'] or '')[:60]} |")

    w("\n## 6. 수치 조건 출현 (M-4 판정 대상)\n")
    w("첨부파일에서 이 조건들이 실제로 얼마나 나오는지 — **C-2 우선순위 결정 근거**\n")
    w("| 조건 | 문서 수 | 공고 수 | 공고 비율 |")
    w("|---|---:|---:|---:|")
    for cue in NUMERIC_CUES:
        w(f"| {cue} | {numeric_tally.get(cue, 0)} | {len(numeric_prog.get(cue, ()))} | {pct(len(numeric_prog.get(cue, ())), n_prog)} |")

    w("\n## 7. 「제외 대상」 문구 패턴\n")
    ex_hits = Counter(h for r in docs for h in r["exclusionHits"])
    ex_prog = {r["pblancId"] for r in docs if r["exclusionHits"]}
    w(f"제외 조건 문구가 있는 공고: **{len(ex_prog)} / {n_prog} ({pct(len(ex_prog), n_prog)})**\n")
    if ex_hits:
        w("| 문구 | 출현 |")
        w("|---|---:|")
        for h, c in ex_hits.most_common(12):
            w(f"| {h} | {c} |")

    OUT_MD.parent.mkdir(parents=True, exist_ok=True)
    OUT_MD.write_text("\n".join(L) + "\n", encoding="utf-8")

    print(f"[survey] 공고 {n_prog}건 / 첨부 {n_att}개 / 추출 성공 {len(docs)}개")
    print(f"[survey] 리포트: {OUT_MD}")
    print(f"[survey] 원본:   {OUT_JSON}")


if __name__ == "__main__":
    main()
